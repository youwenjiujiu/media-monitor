from __future__ import annotations

import io
from datetime import date
from collections import defaultdict
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from models import AnalyzedArticle, ClientConfig

DARK_BLUE = RGBColor(0x1F, 0x3A, 0x5F)
POSITIVE_COLOR = RGBColor(0x27, 0xAE, 0x60)
NEGATIVE_COLOR = RGBColor(0xE7, 0x4C, 0x3C)
NEUTRAL_COLOR = RGBColor(0x7F, 0x8C, 0x8D)

SENTIMENT_COLORS = {
    "Positive": POSITIVE_COLOR,
    "Negative": NEGATIVE_COLOR,
    "Neutral": NEUTRAL_COLOR,
}


def _set_font(run, name="Calibri", size=11, bold=False, color=None):
    run.font.name = name
    run.font.size = Pt(size)
    run.bold = bold
    if color:
        run.font.color.rgb = color


def _add_section_heading(doc: Document, text: str):
    p = doc.add_paragraph()
    run = p.add_run(text)
    _set_font(run, size=14, bold=True, color=DARK_BLUE)
    p.space_after = Pt(6)


def _add_subsection_heading(doc: Document, text: str):
    p = doc.add_paragraph()
    run = p.add_run(text)
    _set_font(run, size=12, bold=True, color=DARK_BLUE)
    p.space_after = Pt(4)


def _add_article_entry(doc: Document, article: AnalyzedArticle):
    analysis = article.analysis
    if not analysis:
        return

    # Headline line: bold headline | source | date
    p = doc.add_paragraph()
    headline_run = p.add_run(analysis.headline)
    _set_font(headline_run, bold=True)

    meta_parts = []
    if analysis.source:
        meta_parts.append(analysis.source)
    if analysis.date:
        meta_parts.append(analysis.date)
    if meta_parts:
        meta_run = p.add_run(f"  |  {' | '.join(meta_parts)}")
        _set_font(meta_run, size=10, color=RGBColor(0x66, 0x66, 0x66))

    # Sentiment tag
    sentiment_color = SENTIMENT_COLORS.get(analysis.sentiment, NEUTRAL_COLOR)
    sentiment_run = p.add_run(f"  [{analysis.sentiment}]")
    _set_font(sentiment_run, size=10, bold=True, color=sentiment_color)

    p.space_after = Pt(2)

    # Summary
    summary_p = doc.add_paragraph()
    summary_run = summary_p.add_run(analysis.summary)
    _set_font(summary_run, size=10)
    summary_p.space_after = Pt(2)

    # URL
    url_p = doc.add_paragraph()
    url_run = url_p.add_run(article.url)
    _set_font(url_run, size=9, color=RGBColor(0x29, 0x80, 0xB9))
    url_p.space_after = Pt(10)


def _add_no_coverage_message(doc: Document, text: str = "No media coverage identified this period."):
    p = doc.add_paragraph()
    run = p.add_run(text)
    _set_font(run, size=10, color=RGBColor(0x99, 0x99, 0x99))
    run.italic = True
    p.space_after = Pt(10)


def generate_report(
    articles: list[AnalyzedArticle], config: ClientConfig
) -> io.BytesIO:
    doc = Document()

    # Set default font
    style = doc.styles["Normal"]
    font = style.font
    font.name = "Calibri"
    font.size = Pt(11)

    # Classify articles
    client_articles = []
    competitor_articles = defaultdict(list)
    industry_articles = defaultdict(list)
    failed_articles = []

    for art in articles:
        if art.analysis_error or not art.analysis:
            failed_articles.append(art)
            continue

        cat = art.analysis.category
        if cat == "client":
            client_articles.append(art)
        elif cat == "competitor":
            comp = art.analysis.competitor_name or "Other"
            competitor_articles[comp].append(art)
        else:
            loc = art.analysis.geographic_location or "Global"
            industry_articles[loc].append(art)

    # Title
    today = date.today()
    week_start = today.strftime("%d %B %Y")
    title_p = doc.add_paragraph()
    title_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_run = title_p.add_run(f"{config.client_name} Media Monitoring")
    _set_font(title_run, size=18, bold=True, color=DARK_BLUE)
    title_p.space_after = Pt(2)

    date_p = doc.add_paragraph()
    date_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    date_run = date_p.add_run(f"Week of {week_start}")
    _set_font(date_run, size=12, color=RGBColor(0x66, 0x66, 0x66))
    date_p.space_after = Pt(20)

    # Section 1: Client Coverage
    _add_section_heading(doc, f"1. {config.client_name} Coverage")
    if client_articles:
        for art in client_articles:
            _add_article_entry(doc, art)
    else:
        _add_no_coverage_message(doc)

    # Section 2: Competitor Coverage
    _add_section_heading(doc, "2. Competitor Coverage")
    for comp_name in config.competitors:
        _add_subsection_heading(doc, comp_name)
        comp_arts = competitor_articles.get(comp_name, [])
        if comp_arts:
            for art in comp_arts:
                _add_article_entry(doc, art)
        else:
            _add_no_coverage_message(doc, "No media coverage identified this period.")

    # Handle competitors not in the predefined list
    for comp_name, comp_arts in competitor_articles.items():
        if comp_name not in config.competitors:
            _add_subsection_heading(doc, comp_name)
            for art in comp_arts:
                _add_article_entry(doc, art)

    # Section 3: Industry News
    _add_section_heading(doc, "3. Industry News")
    if industry_articles:
        for location, arts in industry_articles.items():
            _add_subsection_heading(doc, location)
            for art in arts:
                _add_article_entry(doc, art)
    else:
        _add_no_coverage_message(doc)

    # Failed articles note
    if failed_articles:
        _add_section_heading(doc, "Notes")
        p = doc.add_paragraph()
        run = p.add_run(
            f"{len(failed_articles)} article(s) could not be processed:"
        )
        _set_font(run, size=10, color=RGBColor(0x99, 0x99, 0x99))
        for art in failed_articles:
            err_p = doc.add_paragraph()
            err_run = err_p.add_run(
                f"  - {art.url}: {art.analysis_error or art.scraped.scrape_error or 'Unknown error'}"
            )
            _set_font(err_run, size=9, color=RGBColor(0x99, 0x99, 0x99))

    # Save to BytesIO
    buffer = io.BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer
